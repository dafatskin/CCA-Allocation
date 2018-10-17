#####################
#Presentation module#
#####################

import openpyxl

from docx import Document
from docx.shared import Inches
import datetime
import matplotlib.pyplot as plt
import numpy as np

def autolabel(rects, xpos='center'):
    """
    Attach a text label above each bar in *rects*, displaying its height.

    *xpos* indicates which side to place the text w.r.t. the center of
    the bar. It can be one of the following {'center', 'right', 'left'}.
    """

    xpos = xpos.lower()  # normalize the case of the parameter
    ha = {'center': 'center', 'right': 'left', 'left': 'right'}
    offset = {'center': 0.5, 'right': 0.57, 'left': 0.43}  # x_txt = x + w*off

    for rect in rects:
        height = rect.get_height()
        ax.text(rect.get_x() + rect.get_width()*offset[xpos], 1.01*height,
                '{}'.format(height), ha=ha[xpos], va='bottom')
    
def plotchoicegraph(final_dic):
    #Variables#
    choice_no={1:0,2:0,3:0,4:0,5:0,6:0,7:0,8:0,9:0,0:0}#Rank:Number of rank
    choices=[]
    fig=plt.figure()

    #SetUp
    for name in final_dic:
        for CCA, choice, score in name:
            choice_no[choice]+=1

    for choicetype, choice in choice_no:
        choices.append(int(choice))

    ind = np.arange(len(choices))  # the x locations for the groups
    width = 0.35  # the width of the bars

    fig, ax = plt.subplots()
    rects1 = ax.bar(ind - width/2, choices, width,
                color='SkyBlue')

    # Add some text for labels, title and custom x-axis tick labels, etc.
    ax.set_ylabel('Number of choices')
    ax.set_title('Number of choices allocated')
    ax.set_xticks(ind)
    ax.set_xticklabels(('C1', 'C2', 'C3', 'C4', 'C5', 'C6', 'C7', 'C8', 'C9', 'C0'))
    
    autolabel(rects1, "left")

    return fig
    
def plotallocationgraph(final_dic,list_of_CCA):
    #Variables#
    CCA_no={}
    fig=plt.figure()

    #SetUp
    for cca in list_of_CCA:
        CCA_no[cca]=0
        
    for name in final_dic:
        for CCA, choice, score in final_dic:
            if CCA==CCA_no[cca]:
                CCA_no[cca]+=1
    

    for ccatype, cca in CCA_no:
        list.append(int(cca))

    ind = np.arange(len(CCA_no))  # the x locations for the groups
    width = 0.35  # the width of the bars

    fig, ax = plt.subplots()
    rects1 = ax.bar(ind - width/2, choices, width,
                color='SkyBlue')

    # Add some text for labels, title and custom x-axis tick labels, etc.
    ax.set_ylabel('Number of choices')
    ax.set_title('Number of choices allocated')
    ax.set_xticks(ind)
    ax.set_xticklabels(list_of_CCA)
    
    autolabel(rects1, "left")

    return fig

##def plotsuggestiongraph(final_dic)
    
def final_report(final_dic,list_of_CCA):
    now=datetime.datetime.now()
    
    document = Document()

    document.add_heading('Report for CCA Allocation for ' + str(now.year), 0)

    p = document.add_paragraph('This is the final report for the Year One CCA allocation program, detailing ')
    p.add_run('the final product of the code and some graphical presentations of the final allocations')

    document.add_heading('Content Page', level=1)

    document.add_paragraph(
        'Number of people allocated to their choice', style='List Bullet'
    )
    document.add_paragraph(
        'Number of people allocated to each CCA', style='List Bullet'
    )
    document.add_paragraph(
        'Suggestions for future changes', style='List Bullet'
    )

    document.add_page_break()

    document.add_heading('Number of people allocated to their choice', level=1)
    #document.add_picture(plotchoicegraph(final_dic))

    document.add_page_break()

    document.add_heading('Number of people allocated to each CCA', level=1)
    #document.add_picture(plotallocationgraph(final_dic))

    document.add_page_break()

    document.add_heading('Suggestions for future changes', level=1)
    #document.add_picture(plotsuggestiongraph)

    document.save('final_report.docx')

    
    
def p5_CCAxlsx(final_dic, CCA_dic, merit_CCA_dic, CCA, nameClass, config_vars):
    """
    """
 
    wb = openpyxl.Workbook()
    sheet = wb.active
    sheet.title = CCA

    sheet.column_dimensions["A"].width = 24

    headerfont = openpyxl.styles.Font(size=12, bold=True) #headers
    pdic = {"A1":"NRIC","B1":"Class","C1":"Rank"}
    for cell, header in pdic.items():
        sheet[cell].font = headerfont
        sheet[cell] = header

    datafont = openpyxl.styles.Font(size=11) #data
    row = "1"
    for name, data in CCA_dic[CCA].items():
        row = str(int(row) + 1)
        sheet["A"+row].font = datafont
        sheet["A"+row] = name

        sheet["B"+row].font = datafont
        sheet["B"+row] = nameClass[name]
        
        sheet["C"+row].font = datafont
        if data["rank"] == 0:
            sheet["C"+row] = "DSA"
        elif data["rank"] == 10:
            sheet["C"+row] = ""
        else:
            sheet["C"+row] = data["rank"][-1]

    directory = os.path.abspath("mod_presentation.py")[0:-19]
    savepath = directory + "CCA excel sheets\\" + CCA + " excel sheet.xlsx"
    wb.save(savepath)


def p5_meritCCAxlsx(merit_dic, merit_CCA_dic, CCA, nameClass, config_vars):
    """
    """

    wb = openpyxl.Workbook()
    sheet = wb.active
    sheet.title = CCA

    sheet.column_dimensions["A"].width = 24

    headerfont = openpyxl.styles.Font(size=12, bold=True) #headers
    pdic = {"A1":"NRIC","B1":"Class"}
    for cell, header in pdic.items():
        sheet[cell].font = headerfont
        sheet[cell] = header

    datafont = openpyxl.styles.Font(size=11) #data
    row = "1"
    for name in merit_CCA_dic[CCA]:
        row = str(int(row) + 1)
        sheet["A"+row].font = datafont
        sheet["A"+row] = name

        sheet["B"+row].font = datafont
        sheet["B"+row] = nameClass[name]

    directory = os.path.abspath("mod_presentation.py")[0:-19]
    savepath = directory + "CCA excel sheets\\" + CCA + " excel sheet.xlsx"
    wb.save(savepath)
    

def p5_classxlsx(final_dic, CCA_dic, merit_dic, Class, className, config_vars):
    """
    """

    wb = openpyxl.Workbook()
    sheet = wb.active
    sheet.title = Class

    sheet.column_dimensions["A"].width = 24
    sheet.column_dimensions["B"].width = 17
    
    headerfont = openpyxl.styles.Font(size=12, bold=True) #headers
    pdic = {"A1":"NRIC","B1":"ALLOCATED CCA", "C1":"MERIT CCA"}
    for cell, header in pdic.items():
        sheet[cell].font = headerfont
        sheet[cell] = header

    datafont = openpyxl.styles.Font(size=11) #data
    row = "1"
    for name in className[Class]:
        data = final_dic[name]
        row = str(int(row) + 1)
        sheet["A"+row].font = datafont
        sheet["A"+row] = name

        sheet["B"+row].font = datafont
        sheet["B"+row] = final_dic[name]["CCA"]

        sheet["C"+row].font = datafont
        try:
            if merit_dic.dic[name] != None:
                sheet["C"+row] = merit_dic.dic[name]
            else:
                sheet["C"+row] = ""
        except KeyError:
            pass

    directory = os.path.abspath("mod_presentation.py")[0:-19]
    savepath = directory + "Class excel sheets\\" + Class + " excel sheet.xlsx"
    wb.save(savepath)
    

def p5_allocated(final_dic, CCA_dic, CCA_ranking, config_vars):
    """
    """

    wb = openpyxl.load_workbook("unallocated.xlsx")
    
    wb.create_sheet(index=0, title="classlist_new")
    sheet = wb["classlist_new"]
    for row in wb["classlist"][config_vars["data formats"]["classlist"][0]["slice_coords"][0]:config_vars["data formats"]["classlist"][0]["slice_coords"][1]]:
        for cell in row:
            if cell.row == 1:
                sheet[cell.coordinate] = cell.value
            else:
                name = sheet["C" + str(cell.row)].value
                if sheet[cell.column][0].value == "ALLOCATED CCA":
                    sheet[cell.coordinate] = final_dic[name]["CCA"]
                elif sheet[cell.column][0].value == "CHOICE #":
                    sheet[cell.coordinate] = final_dic[name]["rank"]
                elif sheet[cell.column][0].value == "CCA RANK":
                    try:
                        if name in CCA_ranking[final_dic[name]["CCA"]]:
                            try:
                                sheet[cell.coordinate] = CCA_ranking[final_dic[name]["CCA"]][name]["rank"]
                            except KeyError:
                                sheet[cell.coordinate] = CCA_ranking[final_dic[name]["CCA"]][name]["RANK"]
                        else:
                            pass
                    except KeyError:
                        pass
                else:
                    sheet[cell.coordinate] = cell.value

    del wb["classlist"]
    sheet.title = "classlist"
    
    wb.save("allocated1.xlsx") 
